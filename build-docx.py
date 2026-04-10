#!/usr/bin/env python3
"""
Kombiniert ein Manuskript zu einem DOCX (für Kindle Create Import).

Verwendung:
    python3 build-docx.py                        # -> Manuskript.docx (alle Bücher)
    python3 build-docx.py -b "Buch 1"            # -> nur Buch 1
    python3 build-docx.py -o mein_buch.docx      # -> eigener Dateiname
    python3 build-docx.py -t "Mein Roman"        # -> eigener Titel

Abhängigkeit:
    pip install python-docx
"""

import argparse
import os
import re
import glob

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
MANUSKRIPT_DIR = os.path.join(SCRIPT_DIR, "03 - Manuskript")
RESSOURCEN_DIR = os.path.join(SCRIPT_DIR, "08 - Ressourcen")
LEXIKON_FILE = os.path.join(MANUSKRIPT_DIR, "Lexikon.md")

# Cover-Dateien pro Buch
COVER_FILES = {
    "Buch 1": "ashworth-book1.png",
    "Buch 2": "book2.png",
    "Buch 3": "buck3.png",
    "Buch 4": "buch4.png",
    "Buch 5": "buch5.png",
}

# Serienweite Defaults
DEFAULT_AUTHOR = "C. J. Whitmore"
DEFAULT_SERIES = "Die Ashworth Drei"

SCENE_SEPARATOR = "* * *"


def strip_frontmatter(text: str) -> str:
    m = re.match(r"^---\s*\n.*?\n---\s*\n", text, re.DOTALL)
    if m:
        text = text[m.end():]
    return text


def strip_footer_meta(text: str) -> str:
    text = re.sub(r"\n---\s*\n(\*[^\n]*\*\s*\n?)+\s*$", "", text)
    return text


def get_chapter_title(chapter_dir: str) -> str:
    kapitel_file = os.path.join(chapter_dir, "_Kapitel.md")
    if not os.path.exists(kapitel_file):
        return os.path.basename(chapter_dir)
    with open(kapitel_file, "r", encoding="utf-8") as f:
        content = f.read()
    m = re.search(r"^# (.+)$", content, re.MULTILINE)
    if m:
        title = m.group(1).strip()
        title = re.sub(r"^Kapitel\s+\d+\s*[—–-]\s*", "", title)
        return title
    return os.path.basename(chapter_dir)


def get_scene_files(chapter_dir: str) -> list[str]:
    chapter_name = os.path.basename(chapter_dir)
    pattern = os.path.join(chapter_dir, f"{chapter_name}-S*.md")
    return sorted(glob.glob(pattern))


def extract_scene_text(filepath: str) -> str:
    with open(filepath, "r", encoding="utf-8") as f:
        text = f.read()
    text = strip_frontmatter(text)
    text = strip_footer_meta(text)
    text = re.sub(r"^# .+\n+", "", text)
    text = re.sub(r"^\* \* \*\s*\n*", "", text)
    return text.strip()


def get_chapter_dirs(base_dir: str) -> list[str]:
    dirs = []
    for entry in sorted(os.listdir(base_dir)):
        full = os.path.join(base_dir, entry)
        if os.path.isdir(full) and re.match(r"K\d+", entry):
            dirs.append(full)
    return dirs


def find_cover_image(book_filter: str | None) -> str | None:
    if book_filter and book_filter in COVER_FILES:
        path = os.path.join(RESSOURCEN_DIR, COVER_FILES[book_filter])
        if os.path.exists(path):
            return path
    for name, filename in COVER_FILES.items():
        if book_filter and name != book_filter:
            continue
        path = os.path.join(RESSOURCEN_DIR, filename)
        if os.path.exists(path):
            return path
    return None


def get_book_dirs() -> list[str]:
    book_dirs = []
    for entry in sorted(os.listdir(MANUSKRIPT_DIR)):
        full = os.path.join(MANUSKRIPT_DIR, entry)
        if os.path.isdir(full) and entry.startswith("Buch"):
            book_dirs.append(full)
    if not book_dirs:
        chapter_dirs = get_chapter_dirs(MANUSKRIPT_DIR)
        if chapter_dirs:
            book_dirs = [MANUSKRIPT_DIR]
    return book_dirs


def add_inline_formatting(paragraph, text: str):
    """Fügt Runs mit Kursiv/Fett-Formatierung hinzu (Markdown -> DOCX Runs)."""
    # Tokenize: split on bold (**) and italic (*) markers
    # Process bold first, then italic within each segment
    segments = re.split(r"(\*\*[^*]+\*\*)", text)
    for segment in segments:
        if not segment:
            continue
        bold_match = re.match(r"^\*\*(.+)\*\*$", segment)
        if bold_match:
            run = paragraph.add_run(bold_match.group(1))
            run.bold = True
            continue
        # Process italic within non-bold segments
        italic_parts = re.split(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", segment)
        for i, part in enumerate(italic_parts):
            if not part:
                continue
            run = paragraph.add_run(part)
            if i % 2 == 1:  # Odd indices are italic matches
                run.italic = True


def write_prose(doc: Document, text: str, first_paragraph: bool = True):
    """Schreibt Prosatext als formatierte Absätze."""
    paragraphs = re.split(r"\n\n+", text)
    for p_text in paragraphs:
        p_text = p_text.strip()
        if not p_text:
            continue
        style_name = "First Paragraph" if first_paragraph else "Body Text Prose"
        p = doc.add_paragraph(style=style_name)
        add_inline_formatting(p, p_text)
        first_paragraph = False


def add_scene_separator(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = Pt(24)
    run = p.add_run(SCENE_SEPARATOR)


def build_docx(output_path: str, title: str, author: str,
               book_filter: str | None = None):
    doc = Document()

    # Seitenränder
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.54)

    # Standard-Schriftart
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # Body Text — normaler Prosa-Absatz mit Einzug
    body_style = doc.styles.add_style("Body Text Prose", 1)  # 1 = PARAGRAPH
    body_style.base_style = doc.styles["Normal"]
    body_style.font.name = "Times New Roman"
    body_style.font.size = Pt(12)
    body_style.paragraph_format.space_after = Pt(0)
    body_style.paragraph_format.space_before = Pt(0)
    body_style.paragraph_format.line_spacing = Pt(24)
    body_style.paragraph_format.first_line_indent = Cm(1.27)

    # First Paragraph — erster Absatz nach Kapitel/Szene, kein Einzug
    first_style = doc.styles.add_style("First Paragraph", 1)
    first_style.base_style = doc.styles["Normal"]
    first_style.font.name = "Times New Roman"
    first_style.font.size = Pt(12)
    first_style.paragraph_format.space_after = Pt(0)
    first_style.paragraph_format.space_before = Pt(0)
    first_style.paragraph_format.line_spacing = Pt(24)
    first_style.paragraph_format.first_line_indent = Cm(0)

    # Cover-Bild
    cover_path = find_cover_image(book_filter)
    if cover_path:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(cover_path, width=Inches(4.5))
        doc.add_page_break()
        print(f"Cover: {os.path.basename(cover_path)}")

    # Bücher sammeln
    book_dirs = get_book_dirs()
    if not book_dirs:
        print("Keine Kapitel gefunden!")
        return

    if book_filter:
        book_dirs = [d for d in book_dirs if os.path.basename(d) == book_filter]
        if not book_dirs:
            print(f"Buch '{book_filter}' nicht gefunden!")
            return

    multi_book = len(book_dirs) > 1

    for book_dir in book_dirs:
        book_name = os.path.basename(book_dir)

        if multi_book:
            # Buch-Trennseite (nur bei mehreren Büchern)
            for _ in range(6):
                doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(book_name)
            run.bold = True
            run.font.size = Pt(18)
            doc.add_page_break()
            print(f"\n{book_name}:")

        chapter_dirs = get_chapter_dirs(book_dir)
        for chapter_dir in chapter_dirs:
            chapter_title = get_chapter_title(chapter_dir)
            scene_files = get_scene_files(chapter_dir)

            if not scene_files:
                continue

            # Kapitelüberschrift (Heading 1 für Kindle Create Erkennung)
            heading = doc.add_heading(chapter_title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading.paragraph_format.space_after = Pt(36)

            for i, scene_file in enumerate(scene_files):
                if i > 0:
                    add_scene_separator(doc)
                prose = extract_scene_text(scene_file)
                if prose:
                    write_prose(doc, prose, first_paragraph=True)

            print(f"  {chapter_title}: {len(scene_files)} Szenen")

    # Lexikon am Ende
    if os.path.exists(LEXIKON_FILE):
        with open(LEXIKON_FILE, "r", encoding="utf-8") as f:
            lex_text = f.read()
        lex_text = strip_frontmatter(lex_text)
        m = re.match(r"^# (.+)\n+", lex_text)
        if m:
            lex_title = m.group(1).strip()
            lex_text = lex_text[m.end():]
        else:
            lex_title = "Lexikon"

        heading = doc.add_heading(lex_title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for paragraph in re.split(r"\n\n+", lex_text):
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            if paragraph == "* * *":
                add_scene_separator(doc)
                continue
            h2_match = re.match(r"^##\s+(.+)$", paragraph)
            if h2_match:
                h = doc.add_heading(h2_match.group(1), level=2)
                continue
            if paragraph.startswith(">"):
                lines = paragraph.split("\n")
                for line in lines:
                    line = re.sub(r"^>\s*", "", line).strip().strip("*")
                    if not line:
                        continue
                    p = doc.add_paragraph(style="Quote")
                    p.add_run(line)
                continue
            p = doc.add_paragraph()
            add_inline_formatting(p, paragraph)

        print(f"  {lex_title}")

    doc.save(output_path)
    print(f"\nDOCX erstellt: {output_path}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Manuskript zu DOCX (Kindle Create)")
    parser.add_argument("-o", "--output", default=None,
                        help="Ausgabedatei (Standard: aus Titel abgeleitet)")
    parser.add_argument("-t", "--title", default=DEFAULT_SERIES,
                        help="Titel auf der Titelseite")
    parser.add_argument("-a", "--author", default=DEFAULT_AUTHOR,
                        help="Autor auf der Titelseite")
    parser.add_argument("-b", "--book", default=None,
                        help="Nur ein bestimmtes Buch bauen (z.B. 'Buch 1')")
    args = parser.parse_args()

    output = args.output
    if output is None:
        safe_title = args.title
        for old, new in [("ä", "ae"), ("ö", "oe"), ("ü", "ue"), ("ß", "ss"),
                         ("Ä", "Ae"), ("Ö", "Oe"), ("Ü", "Ue")]:
            safe_title = safe_title.replace(old, new)
        safe_title = re.sub(r"[^\w\s-]", "", safe_title).strip()
        safe_title = re.sub(r"[\s]+", "-", safe_title)
        output = f"{safe_title}.docx"
    if not os.path.isabs(output):
        output = os.path.join(SCRIPT_DIR, output)

    build_docx(output, args.title, args.author, args.book)
